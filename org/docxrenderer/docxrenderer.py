# -*- coding: utf-8 -*-
# Copyright 2015 Alex Yang <aleozlx@126.com>

from bisect import bisect, bisect_left
from StringIO import StringIO
from re import compile as R

# = Requirements ===============
from docx import Document
"""
Documentation 
	https://python-docx.readthedocs.org/en/latest
Python Packages
	https://pypi.python.org/pypi/python-docx/0.7.6
Installation 
	https://python-docx.readthedocs.org/en/latest/user/install.html#install
"""
# ==============================

# available symbols irrelevant of IM @#%&-+=/
r_VARIABLE = R(r'@\s*\+\s*(\w+)\s*/')
r_VARBEGIN = R(r'@\s*(\+\s*\w+|\+|)\s*$')
r_VAREND   = R(r'^\s*(\+\s*\w+|\w+|)\s*/')
r_VARPARTS = [
	(R(r'@\s*$'), R(r'\s*'), R(r'^\s*\+\s*(\w+)\s*/')),
	(R(r'@\s*\+\s*$'), R(r'\s*'), R(r'^\s*(\w+)\s*/')),
	(R(r'@\s*\+\s*(\w+)$'), R(r'(\w+)'), R(r'^(\w+)\s*/')),
	(R(r'@\s*\+\s*(\w+)\s*$'), R(r'\s*'), R(r'^\s*/')),
]

find = lambda text: [(str(m.group(1)), m.span()) for m in r_VARIABLE.finditer(text)]
text = lambda doc: u''.join(p.text for p in doc.paragraphs)

def make_index(e, f=None):
	pi = [[len(p.text),f(p)] for p in e] if f else [[len(p.text),p] for p in e]
	ii = [0] * len(pi)
	for i in xrange(1,len(ii)): ii[i] = ii[i-1] + pi[i-1][0]
	return \
		lambda (s,e): xrange(bisect(ii,s)-1,bisect_left(ii,e)), \
		lambda i: pi[i][1]

def substitute(doc, bindings):
	paragraph_range, get_paragraph = make_index(doc.paragraphs, lambda x:StableParagraph(x))
	for var, span in find(text(doc)):
		if var in bindings:
			pr = paragraph_range(span)
			if len(pr)>1:
				get_paragraph(pr[0]).sub_s(bindings[var])
				get_paragraph(pr[-1]).sub_e()
				for i in list(pr)[1:-1]: get_paragraph(i).truncate()
			elif len(pr)==1:
				get_paragraph(pr[0]).sub(var,bindings[var])
	for table in doc.tables:
		for row in table.rows:
			for cell in row.cells:
				substitute(cell, bindings)

class lazy(object):
	def __init__(self, f):
		super(lazy, self).__init__()
		self.f = f

	def __get__(self, instance, cls):
		val = self.f(instance)
		setattr(instance, self.f.__name__, val)
		return val

class StableParagraph(object):
	def __init__(self, o):
		super(StableParagraph, self).__init__()
		self.o = o

	@lazy
	def index(self): 
		return make_index(self.o.runs)

	@lazy
	def var_spans(self): 
		return find(self.o.text)

	def truncate(self):
		self.o.text=''

	def sub_s(self, val):
		for i in reversed(self.o.runs):
			if r_VARBEGIN.search(i.text):
				i.text = r_VARBEGIN.sub(val, i.text)
				break
			else: i.text=''

	def sub_e(self):
		for i in self.o.runs:
			if r_VAREND.search(i.text):
				i.text = r_VAREND.sub('', i.text)
				break
			else: i.text=''

	def sub(self, var, val):
		run_range, get_run = self.index
		for v, span in self.var_spans:
			if v==var:
				rr = run_range(span)
				first = get_run(rr[0])
				if len(rr)>1:
					first.text = r_VARBEGIN.sub(val, first.text)
					last = get_run(rr[-1])
					last.text = r_VAREND.sub('', last.text)
					for i in list(rr)[1:-1]: get_run(i).text=''
				elif len(rr)==1:
					first.text = r_VARIABLE.sub(lambda m: val if m.group(1)==var else m.group(), first.text)

class DocxTemplate(Document):
	def __init__(self, fname=None):
		if fname==None: super(DocxTemplate, self).__init__()
		else: super(DocxTemplate, self).__init__(fname)

	def get_vars(self):
		a=[str(s) for s in r_VARIABLE.findall(text(self))]
		return {var:a.count(var) for var in set(a)}

	def render(self, **bindings):
		s = StringIO()
		self.save(s)
		doc = Document(s)
		substitute(doc, bindings)
		return doc

if __name__ == '__main__':
	template=DocxTemplate('测试文档.docx')
	print template.get_vars()
	doc=template.render(var1=u'[中文Value1]', var2=u'[中文Value2]', var3=u'[中文Value3]')
	doc.save('测试文档_output.docx')
